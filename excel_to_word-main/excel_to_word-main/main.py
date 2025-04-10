import pandas as pd
import docx
import os
import datetime as dt
import numpy as np
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

### FUNCTIONS ###


def input_ids(excel_file):
    """Reads excel file
    Returns first column as df
    """
    input_df = pd.read_excel(excel_file, header=None)
    return input_df.iloc[:, 0].sort_values()


def cell_colour(doc_table, row_idx, col_idx, colour):
    """Colours a cell based on #code

    doc_table - docx table object where cell is
    row_idx, col_idx - index row/column position of cell
    colour - hex colour code as str
    """
    target_cell = doc_table.rows[row_idx].cells[col_idx]._tc
    table_cell_properties = target_cell.get_or_add_tcPr()
    shade_obj = OxmlElement('w:shd')
    shade_obj.set(qn('w:fill'), colour)
    table_cell_properties.append(shade_obj)

def update_cell_text(doc_table, row_idx, col_idx, text):
    """Updates a cell with new text

    doc_table - docx table object where cell is
    row_idx, col_idx - index row/column position of cell
    colour - hex colour code as str
    """
    cell = doc_table.rows[row_idx].cells[col_idx]
    cell.text = text


def banding(colour1, colour2):
    """Adds alternate banding to rows"""
    while True:
        yield (colour1)  # white
        yield (colour2)  # grey


def merge_rows(doc_table, start_row, start_col, end_row, end_col):
    """Merges cells based on start/end row index and start/end col index
    preserves text from first cell

    args:
    table - docx table object
    num_rows_list - list of numbers for merging points
    column - int index of column to merge to
    """
    cell_text = doc_table.cell(start_row, start_col).text
    start_cell = doc_table.cell(start_row, start_col)
    end_cell = doc_table.cell(end_row, end_col)
    start_cell.merge(end_cell)
    doc_table.cell(start_row, start_col).text = cell_text


def get_sheets(excel_file):
    """Returns list of sheets in excel file"""
    return pd.ExcelFile(excel_file).sheet_names


def file_path(path, filename):
    """Joins a path and filename"""
    return os.path.join(path, filename)


def match_ids(input_table, excel_table, sheet):
    """Reads list of ids and matches against target ids
    Returns a truncated df of matched ids only
    """
    if os.path.exists(excel_table):
        df_excel = pd.read_excel(excel_table, sheet_name=sheet, keep_default_na=False, na_values=(""))
        return df_excel[df_excel.iloc[:, 0].isin(input_table)].fillna('')
    else:
        return print('Master excel table does not exist')


def update_table(df_excel, doc_table):
    """Copy over excel sheet/dataframe rows to docx table rows"""
    num_rows = range(len(df_excel))
    num_cols = range(len(df_excel.columns))

    # iterate over rows in df and add a row
    for r in num_rows:
        row = doc_table.add_row().cells
        # iterate over columns in df and populate word doc with corresponding values
        for c in num_cols:
            row[c].text = str(df_excel.values[r, c])


def get_time_now():
    """Gets the date and time in a given format using datetime"""
    return dt.datetime.now().strftime("%Y%m%d_%H%M%S")


def get_next_row_idx(table_obj):
    """Gets index of next row of table"""
    return len(table_obj.rows) + 1


def get_headers(template):
    """Get a dictionary of table numbers and corresponding len of headers"""
    header_dict = {}
    temp_doc = docx.Document(template)
    for count, table in enumerate(temp_doc.tables):
        num_header_rows = len(table.rows)
        header_dict[count] = num_header_rows
    return header_dict


def update_format(doc, header_style, table_style, template):
    """Takes in current docx object
    Updates the formatting for all headers and table rows based on
    styling in the original word template and saves the file again
    """
    headers_dict = get_headers(template)
    for table in headers_dict:
        header_rows = headers_dict[table]
        for count, row in enumerate(doc.tables[table].rows):
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if count < header_rows:
                            try:
                                paragraph.style = header_style
                            except KeyError:
                                paragraph.style = "Normal"
                        else:
                            try:
                                paragraph.style = table_style
                            except KeyError:
                                paragraph.style = "Normal"
    return doc


def create_table_data(file_name):
    """Creates test excel data"""
    writer = pd.ExcelWriter(file_name)
    cols = list('ABC')
    for x in range(1, 4):
        df_random = pd.DataFrame(np.random.randint(0, 100, size=(100, x)), columns=cols[:x])
        df_random.insert(0, 'ID', np.arange(100)) #df_random['ID'] = np.arange(100)
        df_random.to_excel(writer, f'Sheet{x}', index=False)
    writer.close()

def create_input_data(file_name):
    "Creates test input data"
    writer = pd.ExcelWriter(file_name)
    df_random = pd.DataFrame(np.random.randint(0, 100, size=(10, 1)))
    df_random.to_excel(writer, index=False, header=False)
    writer.close()

def create_test_data(table_name='test_table.xlsx', input_name='input_data.xlsx'):
    create_table_data(table_name)
    create_input_data(input_name)

def create_template(input_table, file_name):
    """Creates a template based on sheets and rows in excel file"""
    doc = docx.Document()
    xl = pd.ExcelFile(input_table)
    for sheet in xl.sheet_names:
        df = pd.read_excel(xl, sheet)
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Table Grid'
        row = table.rows[0].cells
        for c in range(len(df.columns)):
            row[c].text = str(df.columns[c])
        doc.add_paragraph()
    doc.save('template.docx')
    return doc


def main(excel_table, input_data, word_template, output_name, header_style="Heading 1", table_style="Normal"):
    """Copies matching IDs from excel table to word template"""

    if not os.path.exists(word_template):
        doc = create_template(excel_table, 'template.docx')
    else:
        doc = docx.Document(word_template)

    sheets = get_sheets(excel_table)
    target_ids = input_ids(input_data)

    for counter, sheet in enumerate(sheets):
        df_matching = match_ids(target_ids, excel_table, sheet)
        doc_table = doc.tables[counter]
        if not df_matching.empty:
            update_table(df_matching, doc_table)

    doc = update_format(doc, header_style, table_style, word_template)

    doc.save(f"{output_name}_{get_time_now()}.docx")


create_test_data()
main("test_table.xlsx", "input_data.xlsx", 'template.docx', 'test')


'''
Other functions below can be called on specific tables
to merge rows, change the colour of a specific cell and update the text
in a specific cell

merge_rows(doc_table, 3, 0, 3, 1)
cell_colour(doc_table, 1, 1, '#00B050')
update_cell_text(doc_table, 0, 0, 'new_text')
'''
